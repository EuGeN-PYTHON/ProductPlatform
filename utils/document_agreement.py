# from docx import *
# from docx.shared import Inches
import docx
import datetime


def get_path_document(request, order, response_order, form, agreement):
    make_document(request, order, response_order, form, agreement)

def load_document():
    pass

def make_document(request, order, response_order, form, agreement):

    document = docx.Document()
    # id = agreement.pk
    id = 1
    supp_comp = "ООО Поставщик"
    cust_comp = "ООО Заказчик"
    # docx_title=f"Agreement№{id}.docx"
    # ---- Cover Letter ----
    # document.add_picture((r'%s/static/images/my-header.png' % (settings.PROJECT_PATH)), width=Inches(4))
    heading = document.add_heading(f'Соглашение №{id}\nПОСТАВКИ ПРОДУКТОВ ПИТАНИЯ', 1)
    heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(datetime.date.today().strftime(''))

    document.add_paragraph()
    document.add_paragraph()
    # document.add_paragraph("%s" % datetime.now().strftime('%B %d, %Y'))

    document.add_paragraph('Dear Sir or Madam:')
    document.add_paragraph('We are pleased to help you with your widgets.')
    document.add_paragraph('Please feel free to contact me for any additional information.')
    document.add_paragraph('I look forward to assisting you in this project.')

    document.add_paragraph()
    document.add_paragraph('Best regards,')
    document.add_paragraph('Acme Specialist 1]')
    document.add_page_break()
    document.save(f"Agreement№{id}.docx")




    # Prepare document for download
    # -----------------------------
    # f = StringIO()
    # document.save(f)
    # length = f.tell()
    # f.seek(0)
    # response = HttpResponse(
    #     f.getvalue(),
    #     content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    # )
    # response['Content-Disposition'] = 'attachment; filename=' + docx_title
    # response['Content-Length'] = length
    # return response

make_document(1,2,3,4,5)